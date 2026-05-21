import streamlit as st
from streamlit_option_menu import option_menu
import google.generativeai as genai
import pandas as pd
import plotly.express as px
import os
import time
from datetime import datetime
from docx import Document
import io
from fpdf import FPDF
import tempfile
from streamlit_gsheets import GSheetsConnection

class PDFReport(FPDF):
    def header(self):
        try:
            self.image('logo.png', 10, 8, 33)
        except:
            pass
        self.set_font('Arial', 'B', 15)
        self.cell(40)
        self.cell(110, 10, 'Laporan Perkembangan Batin Siswa', 0, 1, 'C')
        self.set_font('Arial', 'I', 10)
        self.cell(40)
        self.cell(110, 10, 'Berchmans Spirit Center', 0, 1, 'C')
        self.ln(20)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf(df, user_name):
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(0, 10, f"Tanggal Laporan: {datetime.now().strftime('%d-%m-%Y')}", ln=True)
    pdf.cell(0, 10, f"Diunduh oleh: {user_name}", ln=True)
    pdf.ln(10)

    total_data = len(df)
    jumlah_konsolasi = len(df[df['Status Awal'] == 'Konsolasi'])
    persen_konsolasi = int((jumlah_konsolasi / total_data) * 100) if total_data > 0 else 0
    persen_desolasi = 100 - persen_konsolasi if total_data > 0 else 0

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Ringkasan Statistik", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, f"Total Data: {total_data}", ln=True)
    pdf.cell(0, 10, f"Konsolasi: {persen_konsolasi}%", ln=True)
    pdf.cell(0, 10, f"Desolasi: {persen_desolasi}%", ln=True)
    pdf.ln(10)

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Priority List (Memerlukan Atensi Khusus)", ln=True)
    pdf.set_font("Arial", size=12)

    df_desolasi_terakhir = df[df['Status Awal'] == 'Desolasi'].tail(15).iloc[::-1]
    if not df_desolasi_terakhir.empty:
        for idx, row in df_desolasi_terakhir.iterrows():
            pdf.cell(0, 8, f"- {row['Nama Siswa']} ({row['Unit']} - {row['Kelas']}) pada {row['Tanggal']}", ln=True)
    else:
        pdf.cell(0, 8, "Tidak ada siswa dalam daftar atensi.", ln=True)

    return pdf.output(dest='S')

# 1. SETUP PAGE
st.set_page_config(page_title="Berchmans Spirit Center", page_icon="🕊️", layout="wide", initial_sidebar_state="expanded")

# --- AUTHENTICATION ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    st.title("Login Sistem")
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit = st.form_submit_button("Login")

        if submit:
            try:
                if username in st.secrets["passwords"] and st.secrets["passwords"][username] == password:
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = username
                    st.success("Login berhasil!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Username atau Password salah.")
            except KeyError:
                st.error("Konfigurasi st.secrets['passwords'] tidak ditemukan.")
    st.stop()

# 2. SETUP API & RADAR MODEL OTOMATIS
try:
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    valid_models = []
    for m in genai.list_models():
        if 'generateContent' in m.supported_generation_methods:
            valid_models.append(m.name)
    if valid_models:
        model_name = valid_models[0].replace("models/", "")
        model = genai.GenerativeModel(model_name)
    else:
        st.error("⚠️ Gagal mengakses model analitik. Periksa konfigurasi sistem.")
except Exception as e:
    st.error(f"⚠️ Masalah koneksi: {e}")

# 3. SETUP DATABASE GOOGLE SHEETS
import streamlit as st
from streamlit_gsheets import GSheetsConnection

try:
    # Kita panggil URL langsung dari secrets, kalau gak ada baru pakai string kosong
    gsheet_url = st.secrets.get("spreadsheet_url", "")
    if not gsheet_url:
        st.error("Spreadsheet URL tidak ditemukan di st.secrets!")
        st.stop()
        
    conn = st.connection("gsheets", type=GSheetsConnection, spreadsheet=gsheet_url)
except Exception as e:
    st.error(f"Gagal menghubungkan ke Google Sheets: {e}")
    st.stop()

# 4. TEMA & WARNA 
st.markdown("""
    <style>
    .main { background-color: #FFFFFF; color: #000000; }
    h1, h2, h3 { color: #0085FF; }
    .stButton>button { background-color: #0085FF; color: white; border-radius: 8px; border: none; }
    .stButton>button:hover { background-color: #FFB800; color: #000000; }
    [data-testid="stSidebar"] { background-color: #f8f9fa !important; }
    [data-testid="stSidebarNav"] {display: none;}
    [data-testid="stMetricValue"] { color: #0085FF; font-weight: bold; }
    p, span, div { color: #000000; }
    </style>
    """, unsafe_allow_html=True)

# 5. SIDEBAR NAVIGASI
with st.sidebar:
    st.image('logo.png', use_container_width=True)
    st.markdown('<p style="color:#8ba1b5; font-size:12px; font-weight:700; letter-spacing:1.5px; margin-bottom: 0px; margin-left: 15px;">MAIN MENU</p>', unsafe_allow_html=True)
    menu = option_menu(
        menu_title=None,
        options=["Dashboard", "Data Input Center", "Student Insights", "Staff Tracker", "Database Management", "Data Archive"],
        icons=["grid", "server", "people", "briefcase", "hdd-stack", "archive"],
        default_index=1, 
        styles={
            "container": {"padding": "0!important", "background-color": "transparent"},
            "icon": {"color": "#8ba1b5", "font-size": "18px"},
            "nav-link": {"font-size": "16px", "text-align": "left", "margin":"5px", "color": "#000000", "--hover-color": "#f0f2f6"},
            "nav-link-selected": {"background-color": "#FFB800", "color": "#000000", "font-weight": "bold", "border-radius": "8px"},
        }
    )

# ==========================================
# HALAMAN 1: DASHBOARD (UI/UX UPGRADE)
# ==========================================
if menu == "Dashboard":
    st.markdown("""
<style>
.dash-title { color: #0085FF; font-size: 24px; font-weight: 800; margin-bottom: 0px; padding-bottom: 0px; }
.dash-subtitle { color: #8ba1b5; font-size: 14px; margin-top: 0px; padding-top: 0px; margin-bottom: 20px;}

.kpi-card { background-color: #ffffff; border-radius: 10px; padding: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); height: 160px; display: flex; flex-direction: column; justify-content: center; margin-bottom: 5px;}
.kpi-card-blue { border: 2px solid #FFB800; }
.kpi-card-yellow { border: 2px solid #FFB800; }
.kpi-card-cyan { border: 2px solid #FFB800; }

.kpi-label { font-size: 12px; font-weight: 700; color: #000000; letter-spacing: 1px; text-transform: uppercase; margin-bottom: 10px; }
.kpi-val { font-size: 24px; font-weight: 800; color: #0085FF; margin-bottom: 5px; line-height: 1.2;}
.kpi-val-yellow { color: #0085FF; }
.kpi-val-cyan { color: #0085FF; }
.kpi-desc { font-size: 13px; color: #8ba1b5; }

.section-container { background-color: #ffffff; border-radius: 12px; padding: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); margin-bottom: 20px; }
.section-title { color: #0085FF; font-size: 18px; font-weight: 800; margin-bottom: 5px; }
.section-subtitle { color: #8ba1b5; font-size: 13px; margin-bottom: 20px; }

.priority-item { border-bottom: 1px solid #f0f2f6; padding-bottom: 10px; margin-bottom: 10px; }
.priority-item:last-child { border-bottom: none; }
.p-name { font-weight: 700; color: #0085FF; font-size: 15px; margin-bottom: 2px; }
.p-class { font-size: 12px; color: #8ba1b5; }
.badge-red { background-color: #ff767533; color: #d63031; padding: 3px 8px; border-radius: 12px; font-size: 10px; font-weight: 800; float: right; }
.p-status { font-size: 12px; color: #d63031; font-weight: 600; margin-top: 5px;}
.p-time { font-size: 11px; color: #b2bec3; float: right; font-style: italic; margin-top: 5px;}
.view-all-btn { display: block; width: 100%; text-align: center; padding: 10px; border: 1px solid #e0e0e0; border-radius: 8px; color: #002244; font-weight: 600; font-size: 13px; text-decoration: none; margin-top: 15px; }
.view-all-btn:hover { background-color: #f4f6f9; }
</style>
    """, unsafe_allow_html=True)

    # --- HEADER ---
    col_header1, col_header2 = st.columns([2.5, 1])
    with col_header1:
        st.markdown('<p class="dash-title" style="color:#0085FF; font-size:28px;">Berchmans Spirit Center</p>', unsafe_allow_html=True)
        st.markdown('<p class="dash-subtitle">Analitik Pergerakan Batin & Evaluasi Pastoral</p>', unsafe_allow_html=True)
    with col_header2:
        st.markdown("<div style='margin-top: 5px;'></div>", unsafe_allow_html=True)
        st.text_input("Search", placeholder="🔍 Cari siswa...", label_visibility="collapsed")

    try:
        df_all = conn.read(worksheet="Data Refleksi", ttl=0)
        df = df_all[df_all['Periode Arsip'] == 'Aktif']
        if not df.empty:
            # --- EXPORT REPORTS ---
            col_ex1, col_ex2, col_ex3 = st.columns([6, 2, 2])
            with col_ex2:
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Dashboard_Data')
                st.download_button(
                    label="📊 Export Excel",
                    data=excel_buffer.getvalue(),
                    file_name=f"Dashboard_Report_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            with col_ex3:
                pdf_bytes = generate_pdf(df, st.session_state.get('username', 'Admin'))
                st.download_button(
                    label="📄 Export PDF Summary",
                    data=pdf_bytes,
                    file_name=f"Summary_Batin_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

            total_data = len(df)
            jumlah_konsolasi = len(df[df['Status Awal'] == 'Konsolasi'])
            persen_konsolasi = int((jumlah_konsolasi / total_data) * 100) if total_data > 0 else 0
            
            df_desolasi = df[df['Status Awal'] == 'Desolasi']

            # --- LOGIKA UNIT ALERT (THRESHOLD) ---
            unit_alert = "Semua Unit Stabil"
            pesan_alert = "Tidak ada tren negatif signifikan"
            
            df_mood_unit = df.groupby('Unit')['Status Awal'].value_counts().unstack(fill_value=0)
            if 'Desolasi' not in df_mood_unit.columns:
                df_mood_unit['Desolasi'] = 0
            if 'Konsolasi' not in df_mood_unit.columns:
                df_mood_unit['Konsolasi'] = 0
                
            df_mood_unit['Total'] = df_mood_unit['Konsolasi'] + df_mood_unit['Desolasi']
            df_mood_unit['Persen_Desolasi'] = (df_mood_unit['Desolasi'] / df_mood_unit['Total']) * 100
            
            df_kritis = df_mood_unit[(df_mood_unit['Desolasi'] >= 3) & (df_mood_unit['Persen_Desolasi'] >= 30)]
            
            if not df_kritis.empty:
                unit_terparah = df_kritis['Persen_Desolasi'].idxmax()
                persen_parah = int(df_kritis.loc[unit_terparah, 'Persen_Desolasi'])
                jumlah_parah = df_kritis.loc[unit_terparah, 'Desolasi']
                
                unit_alert = f"{unit_terparah}: Perlu Atensi!"
                pesan_alert = f"⚠️ {persen_parah}% ({jumlah_parah} laporan) mengalami Desolasi"
            # ----------------------------------------

            # --- ROW 1: KPI CARDS DENGAN EXPANDER ---
            col_kpi1, col_kpi2, col_kpi3 = st.columns(3)
            
            with col_kpi1:
                st.markdown(f"""
<div class="kpi-card kpi-card-blue">
    <div class="kpi-label">SCHOOL-WIDE MOOD 📈</div>
    <div class="kpi-val">{persen_konsolasi}% Konsolasi</div>
    <div class="kpi-desc">Berdasarkan {total_data} total laporan batin</div>
</div>
                """, unsafe_allow_html=True)
                with st.expander("📊 Lihat Detail per Unit"):
                    df_mood = df.groupby(['Unit', 'Status Awal']).size().unstack(fill_value=0).reset_index()
                    if 'Konsolasi' not in df_mood.columns: df_mood['Konsolasi'] = 0
                    if 'Desolasi' not in df_mood.columns: df_mood['Desolasi'] = 0
                    df_mood['Total'] = df_mood['Konsolasi'] + df_mood['Desolasi']
                    df_mood['% Konsolasi'] = (df_mood['Konsolasi'] / df_mood['Total'] * 100).round(1)
                    st.dataframe(df_mood[['Unit', '% Konsolasi', 'Total']], hide_index=True, use_container_width=True)
                
            with col_kpi2:
                st.markdown(f"""
<div class="kpi-card kpi-card-yellow">
    <div class="kpi-label">UNIT ALERT ⚠️</div>
    <div class="kpi-val kpi-val-yellow">{unit_alert}</div>
    <div class="kpi-desc">{pesan_alert}</div>
</div>
                """, unsafe_allow_html=True)
                with st.expander("⚠️ Pantau Unit Lainnya"):
                    if not df_desolasi.empty:
                        df_desolasi_counts = df_desolasi['Unit'].value_counts().reset_index()
                        df_desolasi_counts.columns = ['Unit', 'Jumlah Desolasi']
                        st.dataframe(df_desolasi_counts, hide_index=True, use_container_width=True)
                    else:
                        st.success("Tidak ada data desolasi saat ini.")
                
            with col_kpi3:
                st.markdown("""
<div class="kpi-card kpi-card-cyan">
    <div class="kpi-label">PATTERN DETECTED 🧠</div>
    <div class="kpi-val kpi-val-cyan">Perlu Pemindaian</div>
    <div class="kpi-desc">Klik dropdown di bawah untuk memindai pola batin.</div>
</div>
                """, unsafe_allow_html=True)
                with st.expander("🔍 Pindai Pola Batin"):
                    opsi_pola = ["Seluruh Sekolah"] + sorted(df['Unit'].unique().tolist())
                    target_pola = st.selectbox("Pilih Target Pindaian:", opsi_pola, label_visibility="collapsed")
                    
                    if st.button("Jalankan Pemindaian Pola"):
                        with st.spinner("Memproses data teks..."):
                            df_target = df if target_pola == "Seluruh Sekolah" else df[df['Unit'] == target_pola]
                            teks_refleksi = " ".join(df_target.tail(30)['Refleksi'].dropna().astype(str).tolist())
                            
                            if teks_refleksi.strip():
                                prompt_pola = f"Sebutkan 1 frasa singkat (maksimal 4 kata) yang menjadi tema utama atau masalah paling sering muncul dari kumpulan curhatan ini: '{teks_refleksi}'. Lalu berikan 2 kalimat penjelasan singkat."
                                try:
                                    response = model.generate_content(prompt_pola)
                                    st.success("Pemindaian Selesai!")
                                    st.info(response.text)
                                except Exception as e:
                                    st.error("Gagal terhubung ke AI.")
                            else:
                                st.warning("Data refleksi tidak cukup untuk dipindai.")
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # --- ROW 2: CHART & PRIORITY LIST ---
            col_mid1, col_mid2 = st.columns([2.5, 1.2])
            
            with col_mid1:
                st.markdown("""
<div style="background-color: #ffffff; border-radius: 12px; padding: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); margin-bottom: 10px;">
    <div style="color: #0085FF; font-size: 18px; font-weight: 800; margin-bottom: 5px;">Trend Batin Per Unit</div>
    <div style="color: #8ba1b5; font-size: 13px;">Perbandingan Tingkat Konsolasi vs Desolasi (Data Real-Time)</div>
</div>
                """, unsafe_allow_html=True)
                
                df_unit = df.groupby(['Unit', 'Status Awal']).size().reset_index(name='Jumlah')
                fig_bar = px.bar(
                    df_unit, 
                    y='Unit', 
                    x='Jumlah', 
                    color='Status Awal', 
                    orientation='h', 
                    barmode='group',
                    color_discrete_map={'Konsolasi':'#0085FF', 'Desolasi':'#FFB800'}
                )
                fig_bar.update_layout(
                    margin=dict(t=10, b=0, l=0, r=0),
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                    yaxis_title=None,
                    xaxis_title=None,
                    legend_title=None,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                )
                st.plotly_chart(fig_bar, use_container_width=True)

            with col_mid2:
                st.markdown("""
<div class="section-container" style="padding-bottom: 5px;">
    <div class="section-title">Priority Cura Personalis</div>
    <div class="section-subtitle">High Risk Intervention List</div>
</div>
                """, unsafe_allow_html=True)
                
                df_desolasi_terakhir = df[df['Status Awal'] == 'Desolasi'].tail(3).iloc[::-1]
                
                if not df_desolasi_terakhir.empty:
                    for idx, row in df_desolasi_terakhir.iterrows():
                        st.markdown(f"""
                        <div class="priority-item">
                            <span class="badge-red">PERLU ATENSI</span>
                            <div class="p-name">{row['Nama Siswa']}</div>
                            <div class="p-class">{row['Unit']} - {row['Kelas']}</div>
                            <div><span class="p-status">Desolasi</span> <span class="p-time">{row['Tanggal']}</span></div>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("Tidak ada siswa dalam radar risiko tinggi saat ini.")
                    
                st.markdown('<a href="#" class="view-all-btn">Lihat Semua di Menu Insight ></a>', unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # --- AI SUMMARY PER CLASS TABLE ---
            st.markdown("""
<div class="section-container" style="padding-bottom: 20px; margin-top: 10px;">
    <div class="section-title">AI Summary per Class</div>
    <div class="section-subtitle">Rekapitulasi Dominasi Batin dan Status Atensi Berdasarkan Kelas</div>
</div>
            """, unsafe_allow_html=True)

            df_kelas = df.groupby(['Kelas', 'Status Awal']).size().unstack(fill_value=0)
            if 'Konsolasi' not in df_kelas.columns: df_kelas['Konsolasi'] = 0
            if 'Desolasi' not in df_kelas.columns: df_kelas['Desolasi'] = 0

            df_kelas['Total'] = df_kelas['Konsolasi'] + df_kelas['Desolasi']
            df_kelas['Persen_Desolasi'] = (df_kelas['Desolasi'] / df_kelas['Total']) * 100
            df_kelas['Persen_Konsolasi'] = (df_kelas['Konsolasi'] / df_kelas['Total']) * 100

            html_table = """
<style>
.class-summary-table { width: 100%; border-collapse: collapse; margin-top: -10px; margin-bottom: 20px; }
.class-summary-table th { background-color: #f4f6f9; color: #8ba1b5; font-size: 13px; text-transform: uppercase; padding: 12px; text-align: left; border-bottom: 2px solid #e0e0e0; }
.class-summary-table td { padding: 12px; border-bottom: 1px solid #f0f2f6; color: #002244; font-size: 14px; font-weight: 500; }
.status-aman { color: #27ae60; font-weight: bold; background-color: #e8f8f5; padding: 5px 10px; border-radius: 6px; font-size: 12px;}
.status-warning { color: #d63031; font-weight: bold; background-color: #fadedf; padding: 5px 10px; border-radius: 6px; font-size: 12px;}
</style>
<table class="class-summary-table">
    <thead>
        <tr>
            <th>Nama Kelas</th>
            <th>Dominasi Batin</th>
            <th>Status Atensi</th>
        </tr>
    </thead>
    <tbody>
"""
            for kelas, row in df_kelas.iterrows():
                if row['Total'] == 0: continue

                if row['Persen_Desolasi'] >= 30:
                    status = '<span class="status-warning">Warning</span>'
                else:
                    status = '<span class="status-aman">Aman</span>'

                if row['Konsolasi'] >= row['Desolasi']:
                    dominasi = f"Konsolasi ({int(row['Persen_Konsolasi'])}%)"
                else:
                    dominasi = f"Desolasi ({int(row['Persen_Desolasi'])}%)"

                html_table += f"<tr><td>{kelas}</td><td>{dominasi}</td><td>{status}</td></tr>"

            html_table += "</tbody></table>"
            st.markdown(html_table, unsafe_allow_html=True)

        else:
            st.info("Belum ada data refleksi yang masuk dalam sistem.")
            
    except Exception as e:
        st.error(f"Gagal memuat Dashboard: {e}")

# ==========================================
# HALAMAN 2: DATA INPUT CENTER 
# ==========================================
elif menu == "Data Input Center":
    st.title("Data Input Center 📥")
    st.write("Pusat pencatatan data refleksi harian. Pilih metode input yang sesuai dengan kebutuhan.")
    st.write("---")
    
    df_master_siswa = conn.read(worksheet="Master Siswa", ttl=0)
    
    tab_manual, tab_excel = st.tabs(["📋 Batch Manual Entry (Per Kelas)", "📊 Bulk Excel Import"])
    
    # --- TAB 1: BATCH MANUAL ENTRY ---
    with tab_manual:
        st.markdown("### 📋 Batch Manual Entry by Class")
        st.markdown("<p style='color:#8ba1b5; font-size:14px; margin-top:-10px;'>Rapid entry for multiple students. Isi data sekaligus dan klik simpan di paling bawah.</p>", unsafe_allow_html=True)
        
        if df_master_siswa.empty:
            st.warning("⚠️ Master Data Siswa kosong. Silakan isi daftar nama di menu 'Database Management' terlebih dahulu.")
        else:
            
            col_tgl, col_unit, col_kelas = st.columns(3)
            with col_tgl:
                tanggal_refleksi = st.date_input("Tanggal Refleksi", datetime.now().date())
            with col_unit:
                list_unit = sorted(df_master_siswa['Unit'].dropna().unique().tolist())
                unit_terpilih = st.selectbox("Pilih Unit", list_unit)
            with col_kelas:
                df_unit_itu = df_master_siswa[df_master_siswa['Unit'] == unit_terpilih]
                list_kelas = [k for k in df_unit_itu['Kelas'].unique() if pd.notna(k) and str(k).strip() != '']
                if list_kelas:
                    kelas_terpilih = st.selectbox("Pilih Kelas", sorted(list_kelas))
                else:
                    kelas_terpilih = "-"
            
            if kelas_terpilih != "-":
                df_kelas = df_unit_itu[df_unit_itu['Kelas'] == kelas_terpilih]
                list_nama = sorted(df_kelas['Nama Siswa'].dropna().tolist())
                
                st.write("---")
                
                col_h1, col_h2, col_h3 = st.columns([1.5, 2.5, 3])
                col_h1.markdown("<span style='font-size:12px; font-weight:800; color:#8ba1b5; letter-spacing:1px;'>NAMA SISWA</span>", unsafe_allow_html=True)
                col_h2.markdown("<span style='font-size:12px; font-weight:800; color:#8ba1b5; letter-spacing:1px;'>DOMINASI BATIN</span>", unsafe_allow_html=True)
                col_h3.markdown("<span style='font-size:12px; font-weight:800; color:#8ba1b5; letter-spacing:1px;'>TEKS REFLEKSI</span>", unsafe_allow_html=True)
                st.markdown("<hr style='margin-top: 5px; margin-bottom: 10px; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
                
                with st.form("batch_form", border=False):
                    input_data = []
                    
                    for nama in list_nama:
                        c1, c2, c3 = st.columns([1.5, 2.5, 3])
                        with c1:
                            st.markdown(f"<div style='padding-top:10px; font-weight:600; color:#002244; font-size:14px;'>{nama}</div>", unsafe_allow_html=True)
                        with c2:
                            batin = st.radio("Batin", ["Lewati", "Konsolasi", "Desolasi"], horizontal=True, key=f"batin_{nama}", label_visibility="collapsed")
                        with c3:
                            refleksi = st.text_input("Refleksi", key=f"ref_{nama}", label_visibility="collapsed", placeholder="Ketik refleksi singkat...")
                            
                        input_data.append({"nama": nama, "batin": batin, "refleksi": refleksi})
                        st.markdown("<div style='margin-bottom: 5px; border-bottom: 1px solid #f0f2f6;'></div>", unsafe_allow_html=True)
                        
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    st.markdown("""
                        <style>
                        [data-testid="stFormSubmitButton"] button {
                            background-color: #002244 !important;
                            color: white !important;
                            border-radius: 8px !important;
                            border: none !important;
                            font-weight: bold !important;
                        }
                        [data-testid="stFormSubmitButton"] button:hover {
                            background-color: #dca235 !important;
                            color: #002244 !important;
                        }
                        </style>
                    """, unsafe_allow_html=True)
                    
                    submit_btn = st.form_submit_button("💾 Simpan Semua Data Kelas", use_container_width=True)
                    
                    if submit_btn:
                        data_to_save = []
                        for item in input_data:
                            if item["batin"] != "Lewati":
                                data_to_save.append({
                                    "Tanggal": tanggal_refleksi.strftime("%Y-%m-%d"),
                                    "Unit": unit_terpilih,
                                    "Kelas": kelas_terpilih,
                                    "Nama Siswa": item["nama"],
                                    "Status Awal": item["batin"],
                                    "Refleksi": item["refleksi"],
                                    "Periode Arsip": "Aktif"
                                })
                        
                        if data_to_save:
                            df_batin = conn.read(worksheet="Data Refleksi", ttl=0)
                            df_baru = pd.DataFrame(data_to_save)
                            df_batin = pd.concat([df_batin, df_baru], ignore_index=True)
                            conn.update(worksheet="Data Refleksi", data=df_batin)
                            st.success(f"✅ {len(data_to_save)} data refleksi siswa dari kelas {kelas_terpilih} berhasil disimpan!")
                        else:
                            st.warning("⚠️ Tidak ada data yang disimpan. Pastikan Anda memilih Konsolasi/Desolasi minimal untuk 1 siswa.")

    # --- TAB 2: BULK EXCEL IMPORT ---
    with tab_excel:
        st.markdown("### 📊 Bulk Excel Import")
        st.info("Upload Template Refleksi offline dalam format CSV/Excel. Cocok untuk input data massal dari Google Form yang diunduh ke Excel.")
        file_refleksi = st.file_uploader("Upload File Refleksi", type=['csv', 'xlsx'], key="bulk")
        
        if file_refleksi:
            if st.button("🚀 Simpan Massal ke Database"):
                try:
                    if file_refleksi.name.endswith('.csv'):
                        df_bulk = pd.read_csv(file_refleksi)
                    else:
                        df_bulk = pd.read_excel(file_refleksi)
                    
                    df_batin = conn.read(worksheet="Data Refleksi", ttl=0)
                    data_baru_list = []
                    
                    for i, row in df_bulk.iterrows():
                        tgl_input = row.get('Tanggal', datetime.now().strftime("%Y-%m-%d"))
                        data_baru_list.append({
                            "Tanggal": tgl_input,
                            "Unit": row.get('Unit', '-'),
                            "Kelas": row.get('Kelas', '-'),
                            "Nama Siswa": row.get('Nama Lengkap', row.get('Nama Siswa', 'Anonim')),
                            "Status Awal": row.get('Dominasi Batin', 'Tidak Diketahui'),
                            "Refleksi": row.get('Teks Refleksi', ''),
                            "Periode Arsip": "Aktif"
                        })
                    
                    df_baru = pd.DataFrame(data_baru_list)
                    df_batin = pd.concat([df_batin, df_baru], ignore_index=True)
                    conn.update(worksheet="Data Refleksi", data=df_batin)
                    st.success(f"✅ {len(df_bulk)} baris data berhasil dibaca dan disimpan!")
                except Exception as e:
                    st.error(f"Gagal memproses file: {e}. Pastikan format kolom sesuai dengan standar sistem.")

# ==========================================
# HALAMAN 3: STUDENT INSIGHTS 
# ==========================================
elif menu == "Student Insights":
    st.title("Student Insights & Weekly Report 🔍")
    st.write("Laporan analitik mendalam untuk memantau perkembangan psikologis dan spiritualitas siswa, baik secara individu maupun komunitas.")
    st.write("---")
    
    try:
        df_batin_all = conn.read(worksheet="Data Refleksi", ttl=0)
        df_batin = df_batin_all[df_batin_all['Periode Arsip'] == 'Aktif']

        if df_batin.empty:
            st.warning("Database refleksi aktif masih kosong.")
        else:
            tab1, tab2 = st.tabs(["👤 Rekap Analisis", "📋 Database Lengkap"])
            
            # TAB 1: REKAP ANALISIS
            with tab1:
                col_filter1, col_filter2, col_filter3 = st.columns(3)
                
                with col_filter1:
                    opsi_unit = ["Semua Unit"] + sorted(df_batin['Unit'].unique().tolist())
                    filter_unit = st.selectbox("Filter Unit", opsi_unit)
                    
                with col_filter2:
                    if filter_unit == "Semua Unit":
                        df_filtered_unit = df_batin
                        opsi_kelas = ["Semua Kelas"]
                    else:
                        df_filtered_unit = df_batin[df_batin['Unit'] == filter_unit]
                        opsi_kelas = ["Semua Kelas"] + sorted(df_filtered_unit['Kelas'].unique().tolist())
                    
                    filter_kelas = st.selectbox("Filter Kelas", opsi_kelas)
                    
                with col_filter3:
                    if filter_kelas == "Semua Kelas":
                        df_filtered_kelas = df_filtered_unit
                        opsi_nama = ["Semua Siswa"]
                    else:
                        df_filtered_kelas = df_filtered_unit[df_filtered_unit['Kelas'] == filter_kelas]
                        opsi_nama = ["Semua Siswa"] + sorted(df_filtered_kelas['Nama Siswa'].unique().tolist())
                        
                    filter_nama = st.selectbox("Pilih Nama Siswa", opsi_nama)
                
                st.write("---")
                
                if filter_nama == "Semua Siswa":
                    df_target = df_filtered_kelas.copy()
                    target_analisis = f"Komunitas ({filter_unit} - {filter_kelas})"
                else:
                    df_target = df_filtered_kelas[df_filtered_kelas['Nama Siswa'] == filter_nama].copy()
                    target_analisis = filter_nama
                
                if df_target.empty:
                    st.info(f"Belum ada data refleksi untuk {target_analisis}.")
                else:
                    df_target_terakhir = df_target.tail(50)

                    col_insight1, col_insight2, col_insight3 = st.columns([6, 2, 2])
                    with col_insight1:
                        st.markdown(f"**Riwayat Refleksi: {target_analisis}**")
                    with col_insight2:
                        excel_buffer_insight = io.BytesIO()
                        with pd.ExcelWriter(excel_buffer_insight, engine='openpyxl') as writer:
                            df_target_terakhir.to_excel(writer, index=False, sheet_name='Insight_Data')
                        st.download_button(
                            label="📊 Export Excel",
                            data=excel_buffer_insight.getvalue(),
                            file_name=f"Insight_{target_analisis.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    with col_insight3:
                        pdf_bytes_insight = generate_pdf(df_target_terakhir, st.session_state.get('username', 'Admin'))
                        st.download_button(
                            label="📄 Export PDF Summary",
                            data=pdf_bytes_insight,
                            file_name=f"Summary_{target_analisis.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    df_target_tampil = df_target_terakhir[['Tanggal', 'Kelas', 'Nama Siswa', 'Status Awal', 'Refleksi']].copy()
                    df_target_tampil.insert(0, 'No.', range(1, len(df_target_tampil) + 1))
                    st.dataframe(df_target_tampil, use_container_width=True, hide_index=True)
                    
                    if st.button(f"🧠 Buat Rekap Analisis untuk {target_analisis}"):
                        with st.spinner('Memproses pola analitik data...'):
                            kumpulan_teks = ""
                            for index, row in df_target_terakhir.iterrows():
                                kumpulan_teks += f"- [{row['Tanggal']}] {row['Nama Siswa']} ({row['Kelas']}) - {row['Status Awal']}: {row['Refleksi']}\n"
                            
                            prompt = f"""
                            Sebagai seorang konselor pendidikan dan psikologi sekolah, tugasmu adalah menganalisis rekap jurnal refleksi dari {target_analisis} berdasarkan data berikut. 
                            PENTING: Jangan perkenalkan dirimu atau berbasa-basi. Langsung berikan analisisnya secara profesional.

                            Data Refleksi:
                            {kumpulan_teks}

                            Tolong berikan laporan dengan format:
                            1. **Pola Emosi Dominan:** (Bagaimana tren emosi mayoritas? Apa tema atau pemicu utamanya?)
                            2. **Kesimpulan Ringkas:** (Kondisi psikologis dan dinamika batin saat ini)
                            3. **Rekomendasi Pendampingan:** (Langkah nyata yang sangat spesifik untuk pimpinan unit / wali kelas / guru BK)
                            """
                            
                            try:
                                response = model.generate_content(prompt)
                                st.success("Laporan analitik selesai dibuat.")
                                st.info(response.text)

                                # --- KODE BARU UNTUK DOWNLOAD WORD ---
                                doc = Document()
                                doc.add_heading('LAPORAN ANALITIK BINA IMAN & PSIKOLOGIS', level=1)
                                doc.add_paragraph(f"Target Analisis: {target_analisis}")
                                doc.add_paragraph(f"Tanggal Cetak: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
                                
                                teks_bersih = response.text.replace('**', '')
                                doc.add_paragraph(teks_bersih)
                                
                                bio = io.BytesIO()
                                doc.save(bio)
                                bio.seek(0)
                                
                                st.download_button(
                                    label="📥 Download Laporan (Word)",
                                    data=bio,
                                    file_name=f"Laporan_Analitik_{target_analisis.replace(' ', '_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                # --------------------------------

                            except Exception as e:
                                st.error(f"Gagal memproses sistem: {e}")
            
            # TAB 2: DATABASE LENGKAP
            with tab2:
                st.markdown("### Seluruh Data Refleksi (Disortir per Unit)")
                df_batin_sorted = df_batin.sort_values(by=['Unit', 'Kelas', 'Tanggal'], ascending=[True, True, False]).reset_index(drop=True)
                df_batin_sorted.insert(0, 'No.', range(1, len(df_batin_sorted) + 1))
                st.dataframe(df_batin_sorted, use_container_width=True, hide_index=True)
                
    except FileNotFoundError:
        st.error("Database belum terbentuk.")

# ==========================================
# HALAMAN BARU: STAFF TRACKER
# ==========================================
elif menu == "Staff Tracker":
    st.title("Teacher & Staff Tracker 🧑‍🏫")
    st.write("Modul khusus pendampingan, konseling, dan evaluasi kesejahteraan (well-being) Guru & Staff.")
    st.write("---")

    df_staff = conn.read(worksheet="Master Guru", ttl=0)

    if df_staff.empty:
        st.warning("⚠️ Belum ada data Guru/Staff di Master Data. Pastikan Anda telah mengunggah data pada menu Database Management.")
    else:
        tab_input, tab_riwayat = st.tabs(["📝 Input Konseling", "🗂️ Riwayat & Laporan Akhir"])

        with tab_input:
            col_i1, col_i2 = st.columns([1, 2.5])
            
            with col_i1:
                list_unit_staff = sorted(df_staff['Unit'].dropna().unique().tolist())
                unit_staff = st.selectbox("Pilih Unit", list_unit_staff, key="unit_staff")

                df_staff_filtered = df_staff[df_staff['Unit'] == unit_staff]
                list_nama_staff = sorted(df_staff_filtered['Nama Guru'].dropna().tolist())
                nama_staff = st.selectbox("Nama Guru/Staff", list_nama_staff, key="nama_staff")

                tanggal_konseling = st.date_input("Tanggal Konseling", datetime.now().date(), key="tgl_staff")

            with col_i2:
                detail_konseling = st.text_area("Detail Konseling / Catatan Pendampingan", height=200, placeholder="Tuliskan keluhan, hasil wawancara, atau masalah yang dihadapi di sini...")

                st.markdown("""
                    <style>
                    [data-testid="baseButton-primary"] { background-color: #002244 !important; color: white !important; border: none !important; border-radius: 8px !important; }
                    [data-testid="baseButton-primary"]:hover { background-color: #dca235 !important; color: #002244 !important; }
                    </style>
                """, unsafe_allow_html=True)

                if st.button("💾 Simpan & Analisis AI", use_container_width=True, type="primary"):
                    if detail_konseling:
                        with st.spinner("AI sedang memproses analisis psikologis organisasional..."):
                            prompt = f"""
                            Sebagai seorang Psikolog Industri/Organisasi dan Konselor Pastoral di lingkungan sekolah, analisislah catatan konseling dari staff/guru berikut ini.
                            PENTING: Langsung berikan analisis profesional tanpa basa-basi perkenalan.

                            Nama: {nama_staff}
                            Catatan Konseling: {detail_konseling}

                            Berikan laporan dengan format:
                            1. **Akar Masalah (Root Cause):** (Identifikasi isu utama secara psikologis/profesional)
                            2. **Kondisi Well-being:** (Evaluasi kesejahteraan mental dan emosional saat ini)
                            3. **Rekomendasi Tindak Lanjut:** (Langkah konkrit untuk Kepala Sekolah / HRD dalam mendampingi)
                            """
                            try:
                                response = model.generate_content(prompt)
                                hasil_ai = response.text

                                df_staff_db = conn.read(worksheet="Data Staff", ttl=0)
                                data_baru = pd.DataFrame([{
                                    "Tanggal": tanggal_konseling.strftime("%Y-%m-%d"),
                                    "Unit": unit_staff,
                                    "Nama Staff": nama_staff,
                                    "Detail Konseling": detail_konseling,
                                    "Analisis AI": hasil_ai,
                                    "Periode Arsip": "Aktif"
                                }])
                                df_staff_db = pd.concat([df_staff_db, data_baru], ignore_index=True)
                                conn.update(worksheet="Data Staff", data=df_staff_db)

                                st.success(f"✅ Data konseling {nama_staff} berhasil disimpan dan dianalisis!")
                                st.info(hasil_ai)

                            except Exception as e:
                                st.error(f"Gagal memproses AI: {e}")
                    else:
                        st.error("Catatan konseling wajib diisi!")

        with tab_riwayat:
            st.markdown("### 🗂️ Riwayat Konseling Staff")
            try:
                df_staff_db_all = conn.read(worksheet="Data Staff", ttl=0)
                df_staff_db = df_staff_db_all[df_staff_db_all['Periode Arsip'] == 'Aktif']
                if df_staff_db.empty:
                    st.info("Belum ada riwayat konseling aktif.")
                else:
                    df_staff_db = df_staff_db.sort_values(by='Tanggal', ascending=False).reset_index(drop=True)
                    df_staff_tampil = df_staff_db[['Tanggal', 'Unit', 'Nama Staff', 'Detail Konseling']].copy()
                    df_staff_tampil.insert(0, 'No.', range(1, len(df_staff_tampil) + 1))
                    st.dataframe(df_staff_tampil, use_container_width=True, hide_index=True)

                    st.markdown("<br><hr>", unsafe_allow_html=True)
                    st.markdown("### 🖨️ Cetak Laporan Konseling (Word)")
                    
                    col_r1, col_r2 = st.columns([2, 1])
                    with col_r1:
                        pilihan_cetak = df_staff_db['Nama Staff'] + " (" + df_staff_db['Tanggal'] + ")"
                        laporan_terpilih = st.selectbox("Pilih sesi konseling yang ingin dicetak:", pilihan_cetak.tolist())

                    with col_r2:
                        st.markdown("<div style='margin-top:28px;'></div>", unsafe_allow_html=True)
                        if st.button("Siapkan Dokumen", use_container_width=True):
                            idx = pilihan_cetak == laporan_terpilih
                            row_data = df_staff_db[idx].iloc[0]

                            doc = Document()
                            doc.add_heading('LAPORAN KONSELING GURU & STAFF', level=1)
                            doc.add_paragraph(f"Nama Staff: {row_data['Nama Staff']}")
                            doc.add_paragraph(f"Unit: {row_data['Unit']}")
                            doc.add_paragraph(f"Tanggal Konseling: {row_data['Tanggal']}")
                            
                            doc.add_heading('Catatan Awal Konselor:', level=2)
                            doc.add_paragraph(row_data['Detail Konseling'])
                            
                            doc.add_heading('Analisis Psikologis & Rekomendasi (AI):', level=2)
                            teks_bersih = str(row_data['Analisis AI']).replace('**', '')
                            doc.add_paragraph(teks_bersih)

                            bio = io.BytesIO()
                            doc.save(bio)
                            bio.seek(0)

                            st.download_button(
                                label="📥 Download Laporan (Word)",
                                data=bio,
                                file_name=f"Laporan_Konseling_Staff_{row_data['Nama Staff'].replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary",
                                use_container_width=True
                            )
            except FileNotFoundError:
                st.error("Database konseling belum terbentuk.")

# ==========================================
# HALAMAN 4: DATABASE MANAGEMENT 
# ==========================================
elif menu == "Database Management":
    st.title("Manajemen Data Induk 🗃️")
    st.write("Pusat pengelolaan direktori civitas akademika. Data yang diunggah akan terintegrasi secara otomatis dengan seluruh modul.")
    
    # --- KOTAK METRIK REKAPITULASI ---
    try:
        df_siswa = conn.read(worksheet="Master Siswa", ttl=0)
        df_guru = conn.read(worksheet="Master Guru", ttl=0)

        jumlah_siswa = len(df_siswa) if not df_siswa.empty else 0
        jumlah_guru = len(df_guru) if not df_guru.empty else 0
        total_all = jumlah_siswa + jumlah_guru

        rekap_siswa = df_siswa['Unit'].value_counts().to_dict() if not df_siswa.empty else {}

        st.markdown("""
        <style>
        .rekap-container { display: flex; gap: 15px; margin-top: 20px; margin-bottom: 25px; flex-wrap: wrap; justify-content: space-between;}
        .rekap-box { background-color: #f8f9fa; border-radius: 12px; padding: 20px 10px; text-align: center; border: 1px solid #e9ecef; flex: 1; min-width: 120px; box-shadow: 0 2px 4px rgba(0,0,0,0.02);}
        .rekap-title { font-size: 11px; font-weight: 800; color: #8ba1b5; letter-spacing: 0.5px; text-transform: uppercase; margin-bottom: 8px; line-height: 1.2;}
        .rekap-val { font-size: 28px; font-weight: 900; color: #002244; }
        </style>
        """, unsafe_allow_html=True)

        html_boxes = f'<div class="rekap-container">'
        html_boxes += f'<div class="rekap-box"><div class="rekap-title">TOTAL TERDAFTAR</div><div class="rekap-val">{total_all}</div></div>'

        for unit, jumlah in rekap_siswa.items():
            html_boxes += f'<div class="rekap-box"><div class="rekap-title">UNIT {unit}</div><div class="rekap-val">{jumlah}</div></div>'

        if jumlah_guru > 0:
            html_boxes += f'<div class="rekap-box"><div class="rekap-title">TEACHER & STAFF</div><div class="rekap-val">{jumlah_guru}</div></div>'
            
        html_boxes += '</div>'

        st.markdown(html_boxes, unsafe_allow_html=True)
    except Exception as e:
        pass
        
    st.write("---")

    tab_siswa, tab_guru = st.tabs(["🎓 Master Siswa", "👨‍🏫 Master Guru/Staff"])

    with tab_siswa:
        # --- UPLOAD MASTER SISWA ---
        col_up1, col_up2 = st.columns([1, 1])
        with col_up1:
            st.info("Pastikan format kolom Excel/CSV memuat: **Nama Siswa**, **Unit**, dan **Kelas**.")
            file_siswa = st.file_uploader("Unggah Dokumen Siswa (CSV/Excel)", type=['csv', 'xlsx'], key="upload_siswa")

        with col_up2:
            if file_siswa:
                try:
                    if file_siswa.name.endswith('.csv'):
                        df_upload_siswa = pd.read_csv(file_siswa)
                    else:
                        df_upload_siswa = pd.read_excel(file_siswa)

                    if 'Nama Siswa' not in df_upload_siswa.columns: df_upload_siswa['Nama Siswa'] = None
                    if 'Unit' not in df_upload_siswa.columns: df_upload_siswa['Unit'] = None
                    if 'Kelas' not in df_upload_siswa.columns: df_upload_siswa['Kelas'] = None
                    df_upload_siswa = df_upload_siswa[['Nama Siswa', 'Unit', 'Kelas']]

                    st.warning("Pilih mode penyimpanan untuk Data Siswa:")
                    col_btn1, col_btn2 = st.columns(2)

                    with col_btn1:
                        if st.button("➕ Tambah Data Siswa", use_container_width=True):
                            df_lama_siswa = conn.read(worksheet="Master Siswa", ttl=0)
                            df_gabung_siswa = pd.concat([df_lama_siswa, df_upload_siswa], ignore_index=True)
                            df_gabung_siswa.drop_duplicates(subset=['Nama Siswa', 'Unit', 'Kelas'], keep='last', inplace=True)
                            df_gabung_siswa = df_gabung_siswa.sort_values(by=['Unit', 'Kelas', 'Nama Siswa']).reset_index(drop=True)
                            conn.update(worksheet="Master Siswa", data=df_gabung_siswa)
                            st.success("✅ Data Siswa berhasil diperbarui!")

                    with col_btn2:
                        if st.button("🔄 Reset Data Siswa", use_container_width=True):
                            df_upload_siswa = df_upload_siswa.sort_values(by=['Unit', 'Kelas', 'Nama Siswa']).reset_index(drop=True)
                            conn.update(worksheet="Master Siswa", data=df_upload_siswa)
                            st.success("✅ Seluruh data siswa diganti dengan dokumen baru!")

                except Exception as e:
                    st.error(f"Gagal membaca dokumen: {e}")

        st.write("---")
        
        # --- UI DIREKTORI MASTER SISWA ---
        st.markdown("### 🗂️ Direktori Data Siswa")
        try:
            df_master_siswa = conn.read(worksheet="Master Siswa", ttl=0)
            if not df_master_siswa.empty:
                col_nav, col_search = st.columns([2.5, 1])
                with col_nav:
                    list_unit = ["SEMUA UNIT"] + sorted([str(u).upper() for u in df_master_siswa['Unit'].dropna().unique().tolist()])
                    pilih_unit = option_menu(
                        menu_title=None,
                        options=list_unit,
                        orientation="horizontal",
                        styles={
                            "container": {"padding": "5px", "background-color": "#ffffff", "border": "1px solid #e0e0e0", "border-radius": "10px", "margin":"0px"},
                            "nav-link": {"font-size": "13px", "font-weight": "bold", "color": "#8ba1b5", "margin":"0px", "padding": "10px"},
                            "nav-link-selected": {"background-color": "#002244", "color": "#ffffff", "border-radius": "8px"},
                        },
                        key="menu_unit_db_siswa"
                    )

                with col_search:
                    st.markdown("<div style='margin-top: 10px;'></div>", unsafe_allow_html=True)
                    search_query = st.text_input("Pencarian Siswa", label_visibility="collapsed", placeholder="🔍 Search student name...", key="search_siswa")

                if pilih_unit == "SEMUA UNIT":
                    df_tampil_siswa = df_master_siswa
                else:
                    df_tampil_siswa = df_master_siswa[df_master_siswa['Unit'].str.upper() == pilih_unit]

                list_kelas = [k for k in df_tampil_siswa['Kelas'].unique() if pd.notna(k) and str(k).strip() != '']
                if list_kelas and pilih_unit != "SEMUA UNIT":
                    col_teks, col_pills = st.columns([1, 6])
                    with col_teks:
                        st.markdown("<p style='font-size:13px; font-weight:800; color:#8ba1b5; margin-top:20px; text-align:right; letter-spacing: 1px;'>FILTER KELAS:</p>", unsafe_allow_html=True)
                    with col_pills:
                        opsi_kelas = ["Semua Kelas"] + sorted(list_kelas)
                        pilih_kelas = option_menu(
                            menu_title=None,
                            options=opsi_kelas,
                            orientation="horizontal",
                            styles={
                                "container": {"padding": "0!important", "background-color": "transparent", "margin-top":"10px"},
                                "nav-link": {"font-size": "13px", "color": "#002244", "background-color": "#ffffff", "border": "1px solid #e0e0e0", "border-radius": "20px", "margin": "0 5px", "padding": "5px 15px"},
                                "nav-link-selected": {"background-color": "#dca235", "color": "#002244", "font-weight": "bold", "border": "none"},
                            },
                            key="menu_kelas_db_siswa"
                        )
                        if pilih_kelas != "Semua Kelas":
                            df_tampil_siswa = df_tampil_siswa[df_tampil_siswa['Kelas'] == pilih_kelas]

                if search_query:
                    df_tampil_siswa = df_tampil_siswa[df_tampil_siswa['Nama Siswa'].str.contains(search_query, case=False, na=False)]

                st.markdown("<br>", unsafe_allow_html=True)
                df_tampil_siswa = df_tampil_siswa.sort_values(by=['Unit', 'Kelas', 'Nama Siswa']).reset_index(drop=True)
                df_tampil_siswa.insert(0, 'No.', range(1, len(df_tampil_siswa) + 1))
                st.dataframe(df_tampil_siswa, use_container_width=True, hide_index=True)
                
                st.write("---")
                col_del1, col_del2, col_del3 = st.columns([1,2,1])
                with col_del2:
                    with st.expander("🗑️ Arsipkan & Kosongkan Data Siswa"):
                        st.warning("Tindakan ini akan mengosongkan Master Siswa dan mengarsipkan data refleksi siswa yang aktif.")
                        periode_siswa = st.text_input("Masukkan Nama Periode Arsip (Contoh: 2026/2027 - Ganjil):", key="periode_siswa")
                        if st.button("Ya, Arsipkan Data Siswa", type="primary", use_container_width=True):
                            if periode_siswa.strip() == "":
                                st.error("Nama periode arsip tidak boleh kosong!")
                            else:
                                # Arsipkan data batin
                                try:
                                    df_b = conn.read(worksheet="Data Refleksi", ttl=0)
                                    df_b.loc[df_b['Periode Arsip'] == 'Aktif', 'Periode Arsip'] = periode_siswa
                                    conn.update(worksheet="Data Refleksi", data=df_b)
                                except:
                                    pass

                                conn.update(worksheet="Master Siswa", data=pd.DataFrame(columns=["Nama Siswa", "Unit", "Kelas"]))
                                st.success(f"Data Siswa berhasil dikosongkan dan direkam dalam arsip '{periode_siswa}'. Muat ulang (refresh) halaman.")
            else:
                st.warning("Data Siswa masih kosong.")
        except Exception as e:
             st.warning("Data Siswa masih kosong.")

    with tab_guru:
        # --- UPLOAD MASTER GURU ---
        col_up1, col_up2 = st.columns([1, 1])
        with col_up1:
            st.info("Pastikan format kolom Excel/CSV memuat: **Nama Guru** dan **Unit**.")
            file_guru = st.file_uploader("Unggah Dokumen Guru (CSV/Excel)", type=['csv', 'xlsx'], key="upload_guru")

        with col_up2:
            if file_guru:
                try:
                    if file_guru.name.endswith('.csv'):
                        df_upload_guru = pd.read_csv(file_guru)
                    else:
                        df_upload_guru = pd.read_excel(file_guru)

                    if 'Nama Guru' not in df_upload_guru.columns: df_upload_guru['Nama Guru'] = None
                    if 'Unit' not in df_upload_guru.columns: df_upload_guru['Unit'] = None
                    df_upload_guru = df_upload_guru[['Nama Guru', 'Unit']]

                    st.warning("Pilih mode penyimpanan untuk Data Guru:")
                    col_btn1, col_btn2 = st.columns(2)

                    with col_btn1:
                        if st.button("➕ Tambah Data Guru", use_container_width=True):
                            df_lama_guru = conn.read(worksheet="Master Guru", ttl=0)
                            df_gabung_guru = pd.concat([df_lama_guru, df_upload_guru], ignore_index=True)
                            df_gabung_guru.drop_duplicates(subset=['Nama Guru', 'Unit'], keep='last', inplace=True)
                            df_gabung_guru = df_gabung_guru.sort_values(by=['Unit', 'Nama Guru']).reset_index(drop=True)
                            conn.update(worksheet="Master Guru", data=df_gabung_guru)
                            st.success("✅ Data Guru berhasil diperbarui!")

                    with col_btn2:
                        if st.button("🔄 Reset Data Guru", use_container_width=True):
                            df_upload_guru = df_upload_guru.sort_values(by=['Unit', 'Nama Guru']).reset_index(drop=True)
                            conn.update(worksheet="Master Guru", data=df_upload_guru)
                            st.success("✅ Seluruh data guru diganti dengan dokumen baru!")

                except Exception as e:
                    st.error(f"Gagal membaca dokumen: {e}")

        st.write("---")

        # --- UI DIREKTORI MASTER GURU ---
        st.markdown("### 🗂️ Direktori Data Guru")
        try:
            df_master_guru = conn.read(worksheet="Master Guru", ttl=0)
            if not df_master_guru.empty:
                col_nav, col_search = st.columns([2.5, 1])
                with col_nav:
                    list_unit = ["SEMUA UNIT"] + sorted([str(u).upper() for u in df_master_guru['Unit'].dropna().unique().tolist()])
                    pilih_unit = option_menu(
                        menu_title=None,
                        options=list_unit,
                        orientation="horizontal",
                        styles={
                            "container": {"padding": "5px", "background-color": "#ffffff", "border": "1px solid #e0e0e0", "border-radius": "10px", "margin":"0px"},
                            "nav-link": {"font-size": "13px", "font-weight": "bold", "color": "#8ba1b5", "margin":"0px", "padding": "10px"},
                            "nav-link-selected": {"background-color": "#002244", "color": "#ffffff", "border-radius": "8px"},
                        },
                        key="menu_unit_db_guru"
                    )

                with col_search:
                    st.markdown("<div style='margin-top: 10px;'></div>", unsafe_allow_html=True)
                    search_query = st.text_input("Pencarian Guru", label_visibility="collapsed", placeholder="🔍 Search teacher name...", key="search_guru")

                if pilih_unit == "SEMUA UNIT":
                    df_tampil_guru = df_master_guru
                else:
                    df_tampil_guru = df_master_guru[df_master_guru['Unit'].str.upper() == pilih_unit]

                if search_query:
                    df_tampil_guru = df_tampil_guru[df_tampil_guru['Nama Guru'].str.contains(search_query, case=False, na=False)]

                st.markdown("<br>", unsafe_allow_html=True)
                df_tampil_guru = df_tampil_guru.sort_values(by=['Unit', 'Nama Guru']).reset_index(drop=True)
                df_tampil_guru.insert(0, 'No.', range(1, len(df_tampil_guru) + 1))
                st.dataframe(df_tampil_guru, use_container_width=True, hide_index=True)

                st.write("---")
                col_del1, col_del2, col_del3 = st.columns([1,2,1])
                with col_del2:
                    with st.expander("🗑️ Arsipkan & Kosongkan Data Guru"):
                        st.warning("Tindakan ini akan mengosongkan Master Guru dan mengarsipkan data konseling staff yang aktif.")
                        periode_guru = st.text_input("Masukkan Nama Periode Arsip (Contoh: 2026/2027 - Ganjil):", key="periode_guru")
                        if st.button("Ya, Arsipkan Data Guru", type="primary", use_container_width=True):
                            if periode_guru.strip() == "":
                                st.error("Nama periode arsip tidak boleh kosong!")
                            else:
                                # Arsipkan data staff
                                try:
                                    df_s = conn.read(worksheet="Data Staff", ttl=0)
                                    df_s.loc[df_s['Periode Arsip'] == 'Aktif', 'Periode Arsip'] = periode_guru
                                    conn.update(worksheet="Data Staff", data=df_s)
                                except:
                                    pass

                                conn.update(worksheet="Master Guru", data=pd.DataFrame(columns=["Nama Guru", "Unit"]))
                                st.success(f"Data Guru berhasil dikosongkan dan direkam dalam arsip '{periode_guru}'. Muat ulang (refresh) halaman.")
            else:
                st.warning("Data Guru masih kosong.")
        except Exception as e:
             st.warning("Data Guru masih kosong.")

# ==========================================
# HALAMAN 6: DATA ARCHIVE
# ==========================================
elif menu == "Data Archive":
    st.title("Data Archive 🗄️")
    st.write("Akses riwayat data refleksi dan konseling yang telah diarsipkan dari periode-periode sebelumnya.")
    st.write("---")

    try:
        df_batin_all = conn.read(worksheet="Data Refleksi", ttl=0)
        df_staff_all = conn.read(worksheet="Data Staff", ttl=0)

        arsip_batin = df_batin_all[df_batin_all['Periode Arsip'] != 'Aktif']['Periode Arsip'].unique().tolist()
        arsip_staff = df_staff_all[df_staff_all['Periode Arsip'] != 'Aktif']['Periode Arsip'].unique().tolist()

        semua_arsip = sorted(list(set(arsip_batin + arsip_staff)), reverse=True)

        if not semua_arsip:
            st.info("Belum ada data yang diarsipkan.")
        else:
            periode_pilih = st.selectbox("Pilih Periode Arsip:", semua_arsip)
            st.write("---")

            tab_arsip_siswa, tab_arsip_guru = st.tabs(["🎓 Arsip Refleksi Siswa", "👨‍🏫 Arsip Konseling Staff"])

            with tab_arsip_siswa:
                st.markdown(f"### Riwayat Refleksi Siswa - Periode: {periode_pilih}")
                df_arsip_siswa = df_batin_all[df_batin_all['Periode Arsip'] == periode_pilih]

                if df_arsip_siswa.empty:
                    st.info(f"Tidak ada data refleksi siswa untuk periode {periode_pilih}.")
                else:
                    col_unit_arsip, col_kelas_arsip = st.columns(2)
                    with col_unit_arsip:
                        unit_arsip = st.selectbox("Filter Unit", ["Semua Unit"] + sorted(df_arsip_siswa['Unit'].dropna().unique().tolist()), key="unit_arsip_s")

                    df_tampil_arsip_siswa = df_arsip_siswa
                    if unit_arsip != "Semua Unit":
                        df_tampil_arsip_siswa = df_arsip_siswa[df_arsip_siswa['Unit'] == unit_arsip]

                    with col_kelas_arsip:
                        list_kelas_arsip = ["Semua Kelas"] + sorted(df_tampil_arsip_siswa['Kelas'].dropna().unique().tolist())
                        kelas_arsip = st.selectbox("Filter Kelas", list_kelas_arsip, key="kelas_arsip_s")

                    if kelas_arsip != "Semua Kelas":
                        df_tampil_arsip_siswa = df_tampil_arsip_siswa[df_tampil_arsip_siswa['Kelas'] == kelas_arsip]

                    df_tampil_arsip_siswa = df_tampil_arsip_siswa.sort_values(by=['Tanggal', 'Nama Siswa'], ascending=[False, True]).reset_index(drop=True)
                    df_tampil_arsip_siswa.insert(0, 'No.', range(1, len(df_tampil_arsip_siswa) + 1))
                    st.dataframe(df_tampil_arsip_siswa, use_container_width=True, hide_index=True)

            with tab_arsip_guru:
                st.markdown(f"### Riwayat Konseling Staff - Periode: {periode_pilih}")
                df_arsip_staff = df_staff_all[df_staff_all['Periode Arsip'] == periode_pilih]

                if df_arsip_staff.empty:
                    st.info(f"Tidak ada data konseling staff untuk periode {periode_pilih}.")
                else:
                    unit_arsip_g = st.selectbox("Filter Unit", ["Semua Unit"] + sorted(df_arsip_staff['Unit'].dropna().unique().tolist()), key="unit_arsip_g")

                    df_tampil_arsip_staff = df_arsip_staff
                    if unit_arsip_g != "Semua Unit":
                        df_tampil_arsip_staff = df_arsip_staff[df_arsip_staff['Unit'] == unit_arsip_g]

                    df_tampil_arsip_staff = df_tampil_arsip_staff.sort_values(by=['Tanggal', 'Nama Staff'], ascending=[False, True]).reset_index(drop=True)
                    df_tampil_arsip_staff.insert(0, 'No.', range(1, len(df_tampil_arsip_staff) + 1))
                    st.dataframe(df_tampil_arsip_staff, use_container_width=True, hide_index=True)

    except Exception as e:
        st.error(f"Gagal memuat arsip: {e}")
